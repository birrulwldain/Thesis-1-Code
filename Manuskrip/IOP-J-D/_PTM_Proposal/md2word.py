import sys
import markdown
from docx import Document
from bs4 import BeautifulSoup

def md_to_docx(md_path, docx_path):
    # Read Markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()
    
    # Convert MD to HTML
    html = markdown.markdown(md_text, extensions=['tables'])
    soup = BeautifulSoup(html, 'html.parser')
    
    # Create Word Doc
    doc = Document()
    
    # Simple parse
    for element in soup:
        if element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'h3':
            doc.add_heading(element.text, level=3)
        elif element.name == 'p':
            p = doc.add_paragraph()
            for child in element.children:
                if child.name == 'strong' or child.name == 'b':
                    p.add_run(child.text).bold = True
                elif child.name == 'em' or child.name == 'i':
                    p.add_run(child.text).italic = True
                elif child.name == 'code':
                    p.add_run(child.text).font.name = 'Courier New'
                elif child.string:
                    p.add_run(child.string)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Number')
        elif element.name == 'table':
            rows = element.find_all('tr')
            if rows:
                cols = len(rows[0].find_all(['th', 'td']))
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = 'Table Grid'
                for i, row in enumerate(rows):
                    cells = row.find_all(['th', 'td'])
                    for j, cell in enumerate(cells):
                        if j < cols:
                            table.cell(i, j).text = cell.text
        elif getattr(element, 'name', None) == 'hr':
            doc.add_paragraph('---')

    doc.save(docx_path)
    print(f"Successfully converted to {docx_path}")

if __name__ == "__main__":
    md_to_docx(sys.argv[1], sys.argv[2])
